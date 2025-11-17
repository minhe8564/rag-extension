from .base import BaseExtractionStrategy
from typing import Dict, Any
from loguru import logger
import os
from app.service.minio_client import ensure_bucket, put_object_bytes

try:
    from docx import Document
except ImportError:
    Document = None
    logger.warning("python-docx not installed. DOCX extraction will not work.")


class Docx(BaseExtractionStrategy):
    """Basic DOCX 추출 전략 - python-docx 사용"""
    
    def extract(self, file_path: str) -> Dict[Any, Any]:
        """DOCX 파일 추출 처리"""
        logger.info(f"[BasicDocs] Extracting DOCX file: {file_path}")

        if Document is None:
            raise ImportError("python-docx is required for DOCX extraction. Install it with: pip install python-docx")

        try:
            # DOCX 문서 열기
            doc = Document(file_path)

            # 전체 텍스트 추출
            paragraphs = []
            # 진행률 콜백 (옵션)
            progress_cb = None
            try:
                progress_cb = self.parameters.get("progress_cb") if isinstance(self.parameters, dict) else None
            except Exception:
                progress_cb = None
            total_units = 0
            try:
                total_units = len(doc.paragraphs)
            except Exception:
                total_units = 0
            processed = 0
            if progress_cb and total_units:
                try:
                    progress_cb(processed, total_units)
                except Exception:
                    pass
            for para in doc.paragraphs:
                if para.text.strip():  # 빈 문단 제외
                    paragraphs.append(para.text)
                processed += 1
                if progress_cb and total_units:
                    try:
                        progress_cb(processed, total_units)
                    except Exception:
                        pass

            # 표에서 텍스트 추출
            tables_text = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        row_cells.append(cell.text.strip())
                    table_rows.append(row_cells)
                tables_text.append(table_rows)

            # 전체 텍스트 합치기
            full_text = "\n".join(paragraphs)
            
            # MinIO 업로드
            try:
                user_id = (self.parameters.get("user_id") or "unknown-user") if isinstance(self.parameters, dict) else "unknown-user"
                file_name = (self.parameters.get("file_name") or "extracted.docx") if isinstance(self.parameters, dict) else "extracted.docx"
                base_name = os.path.splitext(file_name)[0] or "extracted"
                object_name = f"{user_id}/{base_name}.txt"
                bucket = "ingest"
                ensure_bucket(bucket)
                put_object_bytes(bucket, object_name, full_text.encode("utf-8"), content_type="text/plain; charset=utf-8")
                return {"full_text": full_text, "bucket": bucket, "path": object_name}
            except Exception as e:
                logger.warning(f"[BasicDocs] MinIO upload failed: {e}")
                return {"full_text": full_text}
        except Exception as e:
            logger.error(f"[BasicDocs] Error extracting DOCX: {str(e)}")
            raise



